"""
Script to create the reference DOCX template with custom styling.

This script should be run once to generate the reference.docx template
used by Pandoc for styling converted documents.

Requirements:
    pip install python-docx

Usage:
    python create_reference_template.py
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os


def create_reference_template():
    """Create a DOCX template with custom styles matching the TipTap editor."""
    doc = Document()

    # Set document margins (1 inch all sides)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Configure Normal style - matches editor prose styling
    normal_style = doc.styles['Normal']
    normal_font = normal_style.font
    normal_font.name = 'Arial'
    normal_font.size = Pt(11)
    normal_font.color.rgb = RGBColor(0, 0, 0)  # Black text
    normal_paragraph = normal_style.paragraph_format
    normal_paragraph.line_spacing = 1.6  # Match editor line height
    normal_paragraph.space_after = Pt(8)  # Add spacing between paragraphs

    # Configure Heading 1 - large, bold, black
    heading1_style = doc.styles['Heading 1']
    heading1_font = heading1_style.font
    heading1_font.name = 'Arial'
    heading1_font.size = Pt(28)  # Larger to match editor prominence
    heading1_font.bold = True
    heading1_font.color.rgb = RGBColor(0, 0, 0)  # Black
    heading1_paragraph = heading1_style.paragraph_format
    heading1_paragraph.space_before = Pt(12)
    heading1_paragraph.space_after = Pt(12)
    heading1_paragraph.line_spacing = 1.2

    # Configure Heading 2 - medium, bold, black
    heading2_style = doc.styles['Heading 2']
    heading2_font = heading2_style.font
    heading2_font.name = 'Arial'
    heading2_font.size = Pt(20)  # Adjusted for hierarchy
    heading2_font.bold = True
    heading2_font.color.rgb = RGBColor(0, 0, 0)  # Black
    heading2_paragraph = heading2_style.paragraph_format
    heading2_paragraph.space_before = Pt(10)
    heading2_paragraph.space_after = Pt(10)
    heading2_paragraph.line_spacing = 1.2

    # Configure Heading 3 - smaller, bold, black
    heading3_style = doc.styles['Heading 3']
    heading3_font = heading3_style.font
    heading3_font.name = 'Arial'
    heading3_font.size = Pt(16)  # Adjusted for hierarchy
    heading3_font.bold = True
    heading3_font.color.rgb = RGBColor(0, 0, 0)  # Black
    heading3_paragraph = heading3_style.paragraph_format
    heading3_paragraph.space_before = Pt(8)
    heading3_paragraph.space_after = Pt(8)
    heading3_paragraph.line_spacing = 1.2

    # Add sample content to establish styles
    doc.add_heading('Sample Heading 1', level=1)
    doc.add_paragraph('This is normal text in Arial 11pt with 1.6 line spacing.')

    doc.add_heading('Sample Heading 2', level=2)
    doc.add_paragraph('Another paragraph of normal text.')

    doc.add_heading('Sample Heading 3', level=3)
    doc.add_paragraph('More sample text.')

    # Add a sample table to establish table styling matching editor
    table = doc.add_table(rows=3, cols=4)
    table.style = 'Table Grid'

    # Set table to have borders (important for Pandoc to pick up)
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml import OxmlElement

    # Function to set cell borders
    def set_cell_border(cell, **kwargs):
        """
        Set cell borders
        Usage:
        set_cell_border(
            cell,
            top={"sz": 12, "val": "single", "color": "#D1D5DB"},
            bottom={"sz": 12, "val": "single", "color": "#D1D5DB"},
            start={"sz": 12, "val": "single", "color": "#D1D5DB"},
            end={"sz": 12, "val": "single", "color": "#D1D5DB"},
        )
        """
        tc = cell._element
        tcPr = tc.get_or_add_tcPr()

        # Check for tag existence, if none found, then create one
        tcBorders = tcPr.find(qn('w:tcBorders'))
        if tcBorders is None:
            tcBorders = OxmlElement('w:tcBorders')
            tcPr.append(tcBorders)

        # List of border positions
        for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
            edge_data = kwargs.get(edge)
            if edge_data:
                tag = 'w:{}'.format(edge)

                element = tcBorders.find(qn(tag))
                if element is None:
                    element = OxmlElement(tag)
                    tcBorders.append(element)

                # Set border attributes
                for key in ["sz", "val", "color", "space", "shadow"]:
                    if key in edge_data:
                        element.set(qn('w:{}'.format(key)), str(edge_data[key]))

    def qn(tag):
        """Get qualified name"""
        from docx.oxml.ns import qn as _qn
        return _qn(tag)

    # Apply borders and styling to all cells
    border_config = {
        "sz": 8, # border size in eighths of a point (8 = 1pt)
        "val": "single",  # border style
        "color": "D1D5DB",  # gray-300 color to match editor
    }

    # Style header row
    header_cells = table.rows[0].cells
    for i, cell in enumerate(header_cells):
        # Set borders
        set_cell_border(
            cell,
            top=border_config,
            bottom=border_config,
            start=border_config,
            end=border_config,
        )

        # Set cell background color to light gray (matching editor's bg-gray-50)
        shading_elm = cell._element.get_or_add_tcPr()
        shading = parse_xml(r'<w:shd {} w:fill="F3F4F6"/>'.format(nsdecls('w')))
        shading_elm.append(shading)

        # Set cell padding for spacious look
        tc = cell._element
        tcPr = tc.get_or_add_tcPr()
        tcMar = parse_xml(r'<w:tcMar {}>'
                         r'<w:top w:w="180" w:type="dxa"/>'  # Increased padding
                         r'<w:left w:w="180" w:type="dxa"/>'
                         r'<w:bottom w:w="180" w:type="dxa"/>'
                         r'<w:right w:w="180" w:type="dxa"/>'
                         r'</w:tcMar>'.format(nsdecls('w')))
        tcPr.append(tcMar)

        # Make header text bold and black
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = paragraph.add_run(f'Header {i+1}')
        run.font.bold = True
        run.font.name = 'Arial'
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0, 0, 0)

    # Add sample data rows with borders
    for row_idx in range(1, 3):
        for col_idx, cell in enumerate(table.rows[row_idx].cells):
            # Set borders
            set_cell_border(
                cell,
                top=border_config,
                bottom=border_config,
                start=border_config,
                end=border_config,
            )

            # Set cell padding
            tc = cell._element
            tcPr = tc.get_or_add_tcPr()
            tcMar = parse_xml(r'<w:tcMar {}>'
                             r'<w:top w:w="180" w:type="dxa"/>'
                             r'<w:left w:w="180" w:type="dxa"/>'
                             r'<w:bottom w:w="180" w:type="dxa"/>'
                             r'<w:right w:w="180" w:type="dxa"/>'
                             r'</w:tcMar>'.format(nsdecls('w')))
            tcPr.append(tcMar)

            # Add data
            paragraph = cell.paragraphs[0]
            run = paragraph.add_run(f'Data {row_idx}-{col_idx+1}')
            run.font.name = 'Arial'
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0, 0, 0)

    # Save the template
    output_path = os.path.join(os.path.dirname(__file__), 'templates', 'reference.docx')
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)

    print(f"Reference template created successfully: {output_path}")


# XML parsing helpers for table styling
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls


if __name__ == "__main__":
    create_reference_template()
